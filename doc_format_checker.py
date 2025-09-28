# docx_formatter.py

import os
import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# --- Helper Functions ---

def set_cell_alignment(cell, alignment):
    """Sets the alignment for all paragraphs within a cell."""
    for p in cell.paragraphs:
        p.alignment = alignment

def set_cell_font_style(cell, is_italic=False, is_bold=False):
    """Sets font style for all runs in a cell."""
    for p in cell.paragraphs:
        for run in p.runs:
            run.italic = is_italic
            run.bold = is_bold
            
def get_table_headers(table):
    """Extracts and cleans header text from the first row of a table."""
    if not table.rows:
        return []
    return [cell.text.strip().lower() for cell in table.rows[0].cells]

# --- Core Formatting Logic ---

def format_paragraphs_and_headings(doc):
    """Formats all paragraphs and headings according to specified rules."""
    in_heading_1_section = False
    
    for p in doc.paragraphs:
        style_name = p.style.name
        
        if style_name.startswith('Heading'):
            in_heading_1_section = False
            if style_name == 'Heading 1':
                in_heading_1_section = True
                font_name, font_size = 'Cambria', 28
            elif style_name == 'Heading 2':
                font_name, font_size = 'Cambria', 20
            elif style_name == 'Heading 3':
                font_name, font_size = 'Cambria', 14
            else:
                font_name = None

            if font_name:
                for run in p.runs:
                    run.font.name = font_name
                    rpr = run._r.get_or_add_rPr()
                    rpr.rFonts.set(qn('w:cs'), font_name)
                    rpr.rFonts.set(qn('w:eastAsia'), font_name)
                    run.font.size = Pt(font_size)
        
        else: # This block handles all non-heading paragraphs
            is_special_font = any(run.font.name in ['Wingdings', 'Wingdings 2', 'Symbol'] for run in p.runs)
            
            if not is_special_font:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                pf = p.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                
                for run in p.runs:
                    run.font.name = 'Segoe UI'
                    run.font.size = Pt(10)
                    rpr = run._r.get_or_add_rPr()
                    rpr.rFonts.set(qn('w:cs'), 'Segoe UI')
                    rpr.rFonts.set(qn('w:eastAsia'), 'Segoe UI')

            if in_heading_1_section and not style_name.startswith('Heading'):
                pf = p.paragraph_format
                pf.line_spacing = 1.33
                pf.space_before = Pt(6)
                pf.space_after = Pt(6)

def format_tables(doc):
    """Identifies and formats tables based on their headers or structure."""
    for table in doc.tables:
        headers = get_table_headers(table)
        
        if "claim element" in headers:
            try:
                claim_idx = headers.index("claim element")
                for row in table.rows:
                    for i, cell in enumerate(row.cells):
                        alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if i == claim_idx else WD_ALIGN_PARAGRAPH.CENTER
                        set_cell_alignment(cell, alignment)
            except ValueError: pass
        elif "publication number" in headers and "inpadoc family members" in headers:
            try:
                pub_idx = headers.index("publication number")
                for row in table.rows:
                    for i, cell in enumerate(row.cells):
                        alignment = WD_ALIGN_PARAGRAPH.CENTER if i == pub_idx else WD_ALIGN_PARAGRAPH.JUSTIFY
                        set_cell_alignment(cell, alignment)
            except ValueError: pass
        elif all(h in headers for h in ["publication number", "title", "priority date", "filing date", "publication date", "inventor(s)", "assignee(s)"]):
            for row in table.rows:
                for cell in row.cells:
                    set_cell_alignment(cell, WD_ALIGN_PARAGRAPH.CENTER)
        elif all(h in headers for h in ["#", "title", "publication date", "source", "author(s)"]):
             for row in table.rows:
                for cell in row.cells:
                    set_cell_alignment(cell, WD_ALIGN_PARAGRAPH.CENTER)
        elif all(h in headers for h in ["logic", "key-string", "hits"]):
            try:
                logic_idx, key_idx = headers.index("logic"), headers.index("key-string")
                for r_idx, row in enumerate(table.rows):
                    for c_idx, cell in enumerate(row.cells):
                        if c_idx == logic_idx: set_cell_alignment(cell, WD_ALIGN_PARAGRAPH.LEFT)
                        elif c_idx == key_idx:
                            set_cell_alignment(cell, WD_ALIGN_PARAGRAPH.JUSTIFY)
                            if r_idx > 0: cell.text = cell.text.upper()
                        else: set_cell_alignment(cell, WD_ALIGN_PARAGRAPH.CENTER)
            except ValueError: pass
        elif "search string" in headers:
             for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == headers.index("search string") else WD_ALIGN_PARAGRAPH.CENTER
                    set_cell_alignment(cell, alignment)
        elif headers == ["#", "name"]:
            try:
                name_idx = headers.index("name")
                for r_idx, row in enumerate(table.rows):
                    if r_idx == 0: continue
                    if len(row.cells) > name_idx:
                        cell = row.cells[name_idx]
                        cell.text = cell.text.upper()
                        set_cell_alignment(cell, WD_ALIGN_PARAGRAPH.CENTER)
            except ValueError: pass
        elif headers == ["#", "claim element", "example sections", "analyst’s comment", "potential relevance"]:
            try:
                rel_idx = headers.index("potential relevance")
                for r_idx, row in enumerate(table.rows):
                    for c_idx, cell in enumerate(row.cells):
                        if c_idx == rel_idx:
                            set_cell_alignment(cell, WD_ALIGN_PARAGRAPH.CENTER)
                            if r_idx > 0: set_cell_font_style(cell, is_italic=True)
                        else: set_cell_alignment(cell, WD_ALIGN_PARAGRAPH.JUSTIFY)
            except ValueError: pass
        elif len(table.columns) == 2 and len(table.rows) >= 4:
            expected_labels = ["Publication Date", "Filing Date", "Abstract", "Relevant Text"]
            actual_labels = [table.cell(i, 0).text.strip() for i in range(4)]
            if actual_labels == expected_labels:
                print("Found specific key-value table. Formatting column 2...")
                for i in range(4):
                    target_cell = table.cell(i, 1)
                    for p in target_cell.paragraphs:
                        if 'w:drawing' in p._p.xml or 'w:pict' in p._p.xml:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            for run in p.runs:
                                if run.font.name not in ['Wingdings', 'Wingdings 2', 'Symbol']:
                                    run.font.name = 'Segoe UI'
                                    run.font.size = Pt(10)

def format_images(doc):
    """Centers all inline images in the document."""
    for p in doc.paragraphs:
        if 'w:drawing' in p._p.xml or 'w:pict' in p._p.xml:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def format_headers_and_footers(doc):
    """Updates the date in headers and footers."""
    current_date = datetime.datetime.now().strftime("%B %d, %Y")
    date_placeholder = "[DATE]" 
    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            if header:
                for p in header.paragraphs:
                    if date_placeholder in p.text:
                        p.text = p.text.replace(date_placeholder, current_date)

def update_doc_properties(doc, filepath):
    """Updates document properties based on the filename."""
    filename = os.path.basename(filepath)
    name_without_ext = os.path.splitext(filename)[0]
    props = doc.core_properties
    props.title, props.keywords = name_without_ext, name_without_ext
    props.subject = f"Formatted Document: {name_without_ext}"
    props.author = "Formatting Application"
    props.comments = f"This document was automatically formatted on {datetime.date.today()}."
    props.category = "Formatted Reports"

# --- Main Processing Function ---
def process_document(input_path):
    """Main function to apply all formatting rules to a DOCX file."""
    doc = Document(input_path)
    
    print("Formatting paragraphs and headings...")
    format_paragraphs_and_headings(doc)
    
    print("Formatting tables...")
    format_tables(doc)
    
    print("Centering images...")
    format_images(doc)
    
    print("Updating headers and footers...")
    format_headers_and_footers(doc)
    
    print("Updating document properties...")
    update_doc_properties(doc, input_path)
    
    # TOC functionality has been removed.
    
    base, ext = os.path.splitext(input_path)
    output_path = f"{base}_formatted{ext}"
    doc.save(output_path)
    
    return output_path