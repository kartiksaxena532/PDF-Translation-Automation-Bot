"""
DOCX Formatter - Backend Formatting Engine
docx_formatter.py

Contains all logic for processing and formatting Word documents.
Applies comprehensive formatting rules while preserving document structure.
"""

import os
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
import re


class DocxFormatter:
    """Main formatter class for processing Word documents."""
    
    def __init__(self, file_path):
        """
        Initialize formatter with document path.
        
        Args:
            file_path (str): Path to the input DOCX file
        """
        self.file_path = file_path
        self.doc = None
        self.output_path = None
        
    def format_document(self):
        """
        Main method to format the entire document.
        
        Returns:
            str: Path to the formatted document
        """
        # Load document
        self.doc = Document(self.file_path)
        
        # Apply all formatting rules
        self.format_paragraphs_and_headings()
        self.format_tables()
        self.format_images()
        self.update_doc_properties()
        self.format_headers_and_footers()
        
        # Save formatted document
        self.output_path = self._generate_output_path()
        self.doc.save(self.output_path)
        
        return self.output_path
    
    def _generate_output_path(self):
        """Generate output file path with '_formatted' suffix."""
        path = Path(self.file_path)
        new_name = f"{path.stem}_formatted{path.suffix}"
        return str(path.parent / new_name)
    
    def format_paragraphs_and_headings(self):
        """Format all paragraphs and headings in the document."""
        previous_style = None
        
        for paragraph in self.doc.paragraphs:
            # Skip empty paragraphs
            if not paragraph.text.strip():
                continue
            
            # Reset spacing to ensure clean formatting
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            
            # Check if paragraph is a heading
            if paragraph.style.name.startswith('Heading'):
                self._format_heading(paragraph)
                previous_style = paragraph.style.name
            else:
                # Format as regular paragraph
                self._format_regular_paragraph(paragraph)
                
                # Apply special spacing if follows Heading 1
                if previous_style == 'Heading 1':
                    paragraph.paragraph_format.line_spacing = 1.33
                    paragraph.paragraph_format.space_before = Pt(6)
                    paragraph.paragraph_format.space_after = Pt(6)
                
                previous_style = paragraph.style.name
    
    def _format_heading(self, paragraph):
        """
        Format heading paragraphs with Cambria font and hierarchical sizing.
        
        Args:
            paragraph: The paragraph object to format
        """
        style_name = paragraph.style.name
        
        # Define heading formats
        heading_formats = {
            'Heading 1': {'size': 28, 'bold': True, 'color': RGBColor(0, 0, 0)},
            'Heading 2': {'size': 20, 'bold': True, 'color': RGBColor(0, 0, 0)},
            'Heading 3': {'size': 14, 'bold': True, 'color': RGBColor(0, 0, 0)},
        }
        
        if style_name in heading_formats:
            fmt = heading_formats[style_name]
            
            # Apply formatting to all runs in the heading
            for run in paragraph.runs:
                run.font.name = 'Cambria'
                run.font.size = Pt(fmt['size'])
                run.font.bold = fmt['bold']
                run.font.color.rgb = fmt['color']
                
                # Set font at XML level for better compatibility
                self._set_font_xml(run, 'Cambria')
            
            # Set paragraph alignment
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _format_regular_paragraph(self, paragraph):
        """
        Format regular paragraphs with Segoe UI font and justified alignment.
        
        Args:
            paragraph: The paragraph object to format
        """
        # Set paragraph alignment to justified
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Format each run in the paragraph
        for run in paragraph.runs:
            # Skip if using symbol fonts (Wingdings, Symbol, etc.)
            if run.font.name and 'symbol' in run.font.name.lower():
                continue
            if run.font.name and 'wingdings' in run.font.name.lower():
                continue
            
            # Apply standard formatting
            run.font.name = 'Segoe UI'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Set font at XML level for better compatibility
            self._set_font_xml(run, 'Segoe UI')
    
    def _set_font_xml(self, run, font_name):
        """
        Set font at XML level to ensure it overrides style definitions.
        
        Args:
            run: The run object to modify
            font_name: Name of the font to apply
        """
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" '
                              f'w:hAnsi="{font_name}" w:cs="{font_name}"/>')
            rPr.insert(0, rFonts)
        else:
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)
            rFonts.set(qn('w:cs'), font_name)
    
    def format_tables(self):
        """Format all tables in the document based on their content type."""
        for table in self.doc.tables:
            if not table.rows:
                continue
            
            # Identify table type and apply appropriate formatting
            table_type = self._identify_table_type(table)
            
            if table_type:
                self._apply_table_formatting(table, table_type)
    
    def _identify_table_type(self, table):
        """
        Identify the type of table based on its headers or structure.
        
        Args:
            table: The table object to analyze
            
        Returns:
            str: Table type identifier or None
        """
        if not table.rows:
            return None
        
        # Get first row text (headers)
        headers = []
        for cell in table.rows[0].cells:
            headers.append(cell.text.lower().strip())
        
        # Check for specific table types based on headers
        header_text = ' '.join(headers)
        
        # Type 1: Claim element table
        if 'claim element' in header_text:
            return 'claim_element'
        
        # Type 2: Publication table
        if 'publication number' in header_text:
            return 'publication'
        
        # Type 3: Logic table
        if 'logic' in header_text and 'operation' in header_text:
            return 'logic'
        
        # Type 4: Search string table
        if 'search string' in header_text or 'search query' in header_text:
            return 'search_string'
        
        # Type 5: Prior art table
        if 'prior art' in header_text or 'reference' in header_text:
            return 'prior_art'
        
        # Type 6: Classification table
        if 'classification' in header_text or 'category' in header_text:
            return 'classification'
        
        # Type 7: Key-value table (2 columns, specific structure)
        if len(table.columns) == 2 and len(table.rows) > 2:
            first_col_text = table.rows[0].cells[0].text.lower()
            if any(kw in first_col_text for kw in ['publication date', 'title', 'inventor', 'assignee']):
                return 'key_value'
        
        return 'default'
    
    def _apply_table_formatting(self, table, table_type):
        """
        Apply formatting to table based on its type.
        
        Args:
            table: The table object to format
            table_type: Type identifier for the table
        """
        # Define column alignments for each table type
        alignments = {
            'claim_element': [WD_ALIGN_PARAGRAPH.JUSTIFY, WD_ALIGN_PARAGRAPH.JUSTIFY],
            'publication': [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
            'logic': [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.JUSTIFY],
            'search_string': [WD_ALIGN_PARAGRAPH.JUSTIFY],
            'prior_art': [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.JUSTIFY, WD_ALIGN_PARAGRAPH.CENTER],
            'classification': [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.JUSTIFY],
            'key_value': [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
            'default': [WD_ALIGN_PARAGRAPH.LEFT]
        }
        
        # Get alignment pattern for this table type
        alignment_pattern = alignments.get(table_type, alignments['default'])
        
        # Set table alignment
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Apply formatting to each row
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                # Determine alignment for this column
                if col_idx < len(alignment_pattern):
                    alignment = alignment_pattern[col_idx]
                else:
                    alignment = alignment_pattern[-1]  # Use last pattern for extra columns
                
                # Apply alignment to all paragraphs in cell
                for paragraph in cell.paragraphs:
                    paragraph.alignment = alignment
                    
                    # Format text in cell
                    for run in paragraph.runs:
                        run.font.name = 'Segoe UI'
                        run.font.size = Pt(9)
                        
                        # Bold headers (first row)
                        if row_idx == 0:
                            run.font.bold = True
                            run.font.size = Pt(10)
                
                # Set vertical alignment
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    def format_images(self):
        """Center all inline images in the document."""
        for paragraph in self.doc.paragraphs:
            # Check if paragraph contains an image
            if self._paragraph_has_image(paragraph):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _paragraph_has_image(self, paragraph):
        """
        Check if a paragraph contains an image.
        
        Args:
            paragraph: The paragraph to check
            
        Returns:
            bool: True if paragraph contains an image
        """
        # Check for drawing elements (modern image format)
        drawing_elements = paragraph._element.xpath('.//w:drawing')
        if drawing_elements:
            return True
        
        # Check for picture elements (legacy image format)
        pict_elements = paragraph._element.xpath('.//w:pict')
        if pict_elements:
            return True
        
        # Alternative method: check for graphic data in runs
        for run in paragraph.runs:
            if run._element.xpath('.//a:blip'):
                return True
        
        return False
    
    def update_doc_properties(self):
        """Update document metadata based on filename."""
        filename = Path(self.file_path).stem
        
        # Clean filename for use as title
        title = filename.replace('_', ' ').replace('-', ' ').title()
        
        # Update core properties
        core_props = self.doc.core_properties
        core_props.title = title
        core_props.subject = f"Formatted document: {title}"
        core_props.keywords = f"formatted, {filename}, automated"
        core_props.category = "Formatted Documents"
        core_props.comments = f"Automatically formatted on {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        
        # Update author if not set
        if not core_props.author:
            core_props.author = "DOCX Formatter"
        
        # Update last modified by
        core_props.last_modified_by = "DOCX Formatter"
    
    def format_headers_and_footers(self):
        """Format headers and footers, replacing date placeholders."""
        current_date = datetime.now().strftime('%B %d, %Y')
        
        # Process all sections in the document
        for section in self.doc.sections:
            # Process header
            if section.header:
                self._process_header_footer(section.header, current_date)
            
            # Process footer
            if section.footer:
                self._process_header_footer(section.footer, current_date)
    
    def _process_header_footer(self, header_footer, current_date):
        """
        Process header or footer to replace date placeholders.
        
        Args:
            header_footer: Header or footer object
            current_date: Formatted current date string
        """
        for paragraph in header_footer.paragraphs:
            for run in paragraph.runs:
                # Replace date placeholder
                if '[DATE]' in run.text:
                    run.text = run.text.replace('[DATE]', current_date)
                
                # Apply standard formatting
                run.font.name = 'Segoe UI'
                run.font.size = Pt(10)