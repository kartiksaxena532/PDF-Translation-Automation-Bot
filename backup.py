import sys, os, re, time, subprocess, asyncio, shutil
from urllib.request import urlopen
from urllib.error import URLError
from pathlib import Path
from mimetypes import guess_type
from threading import Thread
from typing import List, Optional
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import img2pdf
from PIL import Image  # kept for build_pdf normalization
from docx.shared import Pt, Inches ,Cm


# NEW: DOCX support
try:
    from docx import Document
except Exception as e:
    raise RuntimeError("python-docx is required. Install with: pip install python-docx") from e

try:
    import fitz  # PyMuPDF (required)
except Exception as e:
    raise RuntimeError("PyMuPDF (fitz) is required. Install with: pip install pymupdf") from e

from playwright.async_api import async_playwright, TimeoutError as PWTimeout
import tkinter as tk
from tkinter import ttk, filedialog, messagebox

# ------------ Configuration (portable & safe) ------------
APP_NAME = "PDF Translator Bot - By Kartik Saxena"

def app_base_dir(drive_letter: str = None) -> Path:
    """
    Allow the user to manually specify their preferred drive for storage.
    User can fill in the drive letter part themselves.
    """
    if sys.platform.startswith("win"):
        drive_letter = drive_letter or "C:"
        base = Path(drive_letter) / "PDF_Translator_Bot"
        base.mkdir(parents=True, exist_ok=True)
        return base

    root = Path(os.environ.get("XDG_DATA_HOME", Path.home() / ".local" / "share"))
    base = root / "PDF_Translator_Bot"
    base.mkdir(parents=True, exist_ok=True)
    return base

def new_run_dir() -> Path:
    """Unique per run; avoids wiping anything outside this isolated sandbox."""
    from datetime import datetime
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S-%f")
    d = app_base_dir() / f"run-{stamp}"
    (d / "raw").mkdir(parents=True, exist_ok=True)
    (d / "translated").mkdir(parents=True, exist_ok=True)
    return d

# Chrome remote debugging attach
DEBUG_PORT = 9222
REMOTE     = f"http://localhost:{DEBUG_PORT}"
CHROME_CANDIDATES = [
    r"C:\Program Files\Google\Chrome\Application\chrome.exe",
    r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
]
EDGE_CANDIDATES = [
    r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
    r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
]
USER_DATA_DIR = Path(os.environ.get("LOCALAPPDATA", r"C:\Users\Public")) / "Chrome" / "PWProfile"

IMG_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}

# ------------ Helpers ------------
def dbg_ready() -> bool:
    try:
        with urlopen(f"{REMOTE}/json/version", timeout=1.5) as r:
            return r.status == 200
    except Exception:
        try:
            with urlopen(f"{REMOTE}/json/version", timeout=1.5) as r:
                return r.status == 200
        except URLError:
            return False

def find_browser_exe() -> str:
    for p in CHROME_CANDIDATES + EDGE_CANDIDATES:
        if Path(p).is_file():
            return p
    raise FileNotFoundError("Chrome/Edge not found. Update CHROME_CANDIDATES/EDGE_CANDIDATES.")

def launch_chrome_if_needed(log):
    if dbg_ready():
        log("Ahhh....shizz here we go again !")
        return
    exe = find_browser_exe()
    USER_DATA_DIR.mkdir(parents=True, exist_ok=True)
    args = [
        exe,
        f"--remote-debugging-port={DEBUG_PORT}",
        f"--user-data-dir={str(USER_DATA_DIR)}",
        "--no-first-run",
        "--no-default-browser-check",
    ]
    creationflags = 0
    if sys.platform.startswith("win"):
        creationflags = subprocess.CREATE_NEW_PROCESS_GROUP | subprocess.DETACHED_PROCESS
    subprocess.Popen(args, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, creationflags=creationflags)
    log("Launching the Dementor(RAM Eater)")
    for _ in range(60):
        if dbg_ready():
            log("Do we have something like Expecto Patronum for chrome also.")
            return
        time.sleep(0.25)
    raise TimeoutError("Could not start Chrome with remote debugging port.")

def wipe_images_only(folder: Path):
    if not folder.exists():
        return
    for p in folder.iterdir():
        if p.is_file() and p.suffix.lower() in IMG_EXTS:
            try:
                p.unlink(missing_ok=True)
            except:
                pass

# ------------ Core rendering: PyMuPDF ONLY ------------
def extract_pages(pdf_path: str, dpi: int, raw_dir: Path, log) -> List[Path]:
    """
    Render every page of the PDF.
    """
    raw_dir.mkdir(parents=True, exist_ok=True)
    # clean previous images
    for p in raw_dir.glob("*"):
        if p.is_file() and p.suffix.lower() in IMG_EXTS:
            try:
                p.unlink(missing_ok=True)
            except:
                pass

    out_paths: List[Path] = []

    # 72 dpi is the base (user space) in PDF; scale to requested dpi
    zoom = max(1.0, dpi / 72.0)
    mat = fitz.Matrix(zoom, zoom)

    with fitz.open(pdf_path) as doc:
        total = doc.page_count
        log(f"Scanning your PDF with → {total} pages → {raw_dir}")
        for i, page in enumerate(doc, 1):
            pix = page.get_pixmap(matrix=mat, alpha=False)  # RGB
            dst = raw_dir / f"page-{i:03}.png"
            pix.save(dst.as_posix())
            out_paths.append(dst)
            log(f"  ✓ {dst.name}")

    if not out_paths:
        raise RuntimeError("No images were produced.Please contact Kartik😭")

    return out_paths

# ------------ Simple DOCX writer (NEW) ------------
class DocxLogger:
    def __init__(self, docx_path: Path, title: Optional[str] = None):
        self.path = Path(docx_path)
        self.doc = Document()

        # Set margins to 2.12 inches on all sides
        section = self.doc.sections[0]    
        section.left_margin = Cm(2.12)
        section.right_margin = Cm(2.12)

        # Set default Normal style
        normal = self.doc.styles["Normal"]
        nf = normal.font
        nf.name = "Segoe UI"
        nf.size = Pt(10)
        # Ensure all scripts use Segoe UI (Windows quirk)
        try:
            rFonts = normal.element.rPr.rFonts
        except AttributeError:
            rFonts = normal.element.get_or_add_rPr().get_or_add_rFonts()
        rFonts.set(qn("w:ascii"), "Segoe UI")
        rFonts.set(qn("w:hAnsi"), "Segoe UI")
        rFonts.set(qn("w:eastAsia"), "Segoe UI")
        rFonts.set(qn("w:cs"), "Segoe UI")

        # Alignment and spacing like screenshot
        pf = normal.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.33

        # Make Heading styles match too (font + size + justify + spacing)
        for h in ("Heading 1", "Heading 2", "Heading 3"):
            if h in self.doc.styles:
                hs = self.doc.styles[h]
                hf = hs.font
                hf.name = "Segoe UI"
                hf.size = Pt(10)
                try:
                    rFonts = hs.element.rPr.rFonts
                except AttributeError:
                    rFonts = hs.element.get_or_add_rPr().get_or_add_rFonts()
                rFonts.set(qn("w:ascii"), "Segoe UI")
                rFonts.set(qn("w:hAnsi"), "Segoe UI")
                rFonts.set(qn("w:eastAsia"), "Segoe UI")
                rFonts.set(qn("w:cs"), "Segoe UI")
                hpf = hs.paragraph_format
                hpf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                hpf.space_before = Pt(0)
                hpf.space_after = Pt(0)
                hpf.line_spacing = 1.33

        if title:
            p = self.doc.add_heading(title, level=1)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.33
            self._force_para_font(p)
            self.doc.add_paragraph("")  # spacer

    def _force_para_font(self, para):
        """Make sure all runs in this paragraph are Segoe UI / 10pt and spacing is correct."""
        for run in para.runs:
            run.font.name = "Segoe UI"
            run.font.size = Pt(10)
            try:
                rFonts = run._element.rPr.rFonts
            except AttributeError:
                rFonts = run._element.get_or_add_rPr().get_or_add_rFonts()
            rFonts.set(qn("w:ascii"), "Segoe UI")
            rFonts.set(qn("w:hAnsi"), "Segoe UI")
            rFonts.set(qn("w:eastAsia"), "Segoe UI")
            rFonts.set(qn("w:cs"), "Segoe UI")

        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.33
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def add_section(self, heading: str, text: str):
        # Heading
        hp = self.doc.add_heading(heading, level=2)
        hp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        hp.paragraph_format.space_before = Pt(0)
        hp.paragraph_format.space_after = Pt(0)
        hp.paragraph_format.line_spacing = 1.33
        self._force_para_font(hp)

        # Body (split into neat paragraphs)
        for line in text.splitlines():
            p = self.doc.add_paragraph(line if line.strip() else "")
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.33
            self._force_para_font(p)

        # spacer
        sp = self.doc.add_paragraph("")
        sp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(0)
        sp.paragraph_format.line_spacing = 1.33
        self._force_para_font(sp)

    def save(self):
        # Final sweep (paranoia): enforce on anything added by other code
        for p in self.doc.paragraphs:
            if p.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.33
            self._force_para_font(p)
        self.doc.save(self.path.as_posix())

# ------------ Translate via Google Translate Images (Playwright) ------------
async def translate_images(
    img_paths: List[Path],
    target_lang: str,
    close_browser: bool,
    trans_dir: Path,
    log=None,
    txt_append_path: Optional[Path] = None,   # legacy arg (ignored if DOCX used)
    docx_logger: Optional[DocxLogger] = None, # NEW: write into DOCX per image
) -> List[Path]:
    """
    Upload images to Google Translate, download translated images,
    and write extracted text to DOCX (one section per image).
    """
    trans_dir.mkdir(parents=True, exist_ok=True)
    for p in trans_dir.glob("*"):
        if p.is_file() and p.suffix.lower() in IMG_EXTS:
            try: p.unlink(missing_ok=True)
            except: pass

    results: List[Path] = []
    url = f"https://translate.google.co.in/?sl=auto&tl={target_lang}&op=images"
    launch_chrome_if_needed(log or (lambda *_: None))

    async with async_playwright() as pw:
        browser = await pw.chromium.connect_over_cdp(REMOTE)
        try:
            ctx = browser.contexts[0] if browser.contexts else await browser.new_context(accept_downloads=True)

            # Grant clipboard permissions so Copy text works reliably
            origin = "https://translate.google.co.in"
            try:
                await ctx.grant_permissions(["clipboard-read", "clipboard-write"], origin=origin)
            except Exception:
                pass

            page = await ctx.new_page()
            await page.goto(url)

            browse_btn   = page.get_by_role("button", name=re.compile(r"Browse your files", re.I))
            download_btn = page.get_by_role("button", name=re.compile(r"(Download translation|Download)", re.I))
            clear_btn    = page.get_by_role("button", name=re.compile(r"(Clear image|Clear)", re.I))
            show_translated = page.get_by_text("Show translated", exact=True)
            copy_btn_role  = page.get_by_role("button", name=re.compile(r"Copy text", re.I))
            copy_btn_css   = page.locator('button[aria-label="Copy text"]')  # fallback

            total = len(img_paths)
            for idx, img in enumerate(img_paths, 1):
                log(f"[{idx}/{total}] {img.name}")

                payload = {
                    "name": img.name,
                    "mimeType": guess_type(img.name)[0] or "application/octet-stream",
                    "buffer": img.read_bytes()
                }
                try:
                    await page.locator('input[type="file"]').set_input_files(payload, timeout=1500)
                except Exception:
                    async with page.expect_file_chooser() as fc:
                        await browse_btn.click()
                    chooser = await fc.value
                    await chooser.set_files(payload)

                if await show_translated.count() > 0:
                    await show_translated.first.click()

                await download_btn.first.wait_for(state="visible", timeout=60000)
                async with page.expect_download() as dl_info:
                    await download_btn.first.click()
                dl = await dl_info.value
                suggested = dl.suggested_filename or f"{img.stem}"
                ext = Path(suggested).suffix or ".png"
                out_img = trans_dir / f"{img.stem}-{ext}"
                await dl.save_as(str(out_img))
                results.append(out_img)

                # Copy text and write to DOCX (preferred)
                copied = ""
                try:
                    target_btn = copy_btn_role if await copy_btn_role.count() else copy_btn_css
                    if await target_btn.count():
                        await target_btn.first.click()
                        await page.wait_for_timeout(150)
                        copied = await page.evaluate("navigator.clipboard.readText()")
                except Exception:
                    copied = ""

                if copied and docx_logger:
                    try:
                        docx_logger.add_section(img.name, copied.strip())
                        log(f"   ↳ I've append {img.name} to your  DOCX")
                    except Exception as e:
                        log(f"⚠ Could not write text for {img.name}: {e}")
                elif copied and txt_append_path:
                    # fallback: keep old .txt behavior if someone passes it
                    try:
                        with open(txt_append_path, "a", encoding="utf-8") as fp:
                            fp.write(f"\n===== {img.name} =====\n")
                            fp.write(copied.strip())
                            fp.write("\n")
                        log(f"   ↳ You serously named it {img.name}? Whatever check {txt_append_path.name} for the output!")
                    except Exception as e:
                        log(f"⚠ Could not append text for {img.name}: {e}")

                # Clear for next image
                try: await clear_btn.click()
                except: pass

            await page.close()
        finally:
            if close_browser:
                try:
                    cdp = await browser.new_browser_cdp_session()
                    await cdp.send("Browser.close")
                except Exception:
                    pass
                try:
                    await browser.close()
                except Exception:
                    pass

    return results

def build_pdf(images: List[Path], out_pdf: Path, log):
    log("Building Your Final PDF…")
    normalized: List[Path] = []
    for p in images:
        if p.suffix.lower() in (".png", ".jpg", ".jpeg"):
            normalized.append(p); continue
        with Image.open(p) as im:
            q = p.with_suffix(".png"); im.save(q); normalized.append(q)
    with open(out_pdf, "wb") as f:
        f.write(img2pdf.convert([str(p) for p in normalized]))
    log(f"✓ Saved → {out_pdf}")

async def translate_pdf(input_pdf: str, output_pdf: str, target_lang: str = "en", dpi: int = 150, close_browser: bool = True, log=print):
    if not Path(input_pdf).is_file():
        raise FileNotFoundError(f"File not found: {input_pdf}")

    run_dir = new_run_dir()
    raw_dir = run_dir / "raw"
    trans_dir = run_dir / "translated"

    # Ensure output directory exists / recover if invalid
    try:
        Path(output_pdf).parent.mkdir(parents=True, exist_ok=True)
    except Exception:
        fallback = app_base_dir() / Path(output_pdf).name
        log(f"⚠ Output folder not available. Using: {fallback}")
        output_pdf = str(fallback)

    # NEW: build a fresh DOCX path next to the output PDF
    docx_path = Path(output_pdf).with_suffix(".docx")
    docx_logger = DocxLogger(docx_path, title="Translated Text")

    pages = extract_pages(input_pdf, dpi, raw_dir, log)
    translated = await translate_images(
        pages, target_lang, close_browser, trans_dir, log,
        txt_append_path=None,        # we’re using DOCX, not TXT
        docx_logger=docx_logger,     # <- write text here
    )
    build_pdf(translated, Path(output_pdf), log)

    # Save DOCX
    try:
        docx_logger.save()
        log(f"✓ Saved DOCX → {docx_path}")
    except Exception as e:
        log(f"⚠ Could not save DOCX: {e}")

    wipe_images_only(raw_dir)
    wipe_images_only(trans_dir)
    log(f"Cleaned up temporary images in {run_dir}.")

# ------------ GUI ------------
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title(APP_NAME + " ✨")
        self.geometry("900x500")
        self.configure(bg="#f4f6f8")   # light soothing background
        self.resizable(False, False)

        # Modern ttk styling
        style = ttk.Style(self)
        style.theme_use("clam")

        # Soft colors & fonts
        style.configure("TLabel", font=("Segoe UI", 11), background="#f4f6f8", foreground="#333")
        style.configure("Card.TFrame", background="#ffffff", relief="flat")
        style.configure("Sidebar.TFrame", background="#e8ebef")
        style.configure("TButton",
                        font=("Segoe UI", 10, "bold"),
                        padding=8,
                        relief="flat",
                        background="#4a90e2",
                        foreground="white")
        style.map("TButton",
                  background=[("active", "#357ab8")],
                  foreground=[("active", "white")])
        style.configure("TEntry", padding=6, relief="flat", fieldbackground="#fff")
        style.configure("TLabelframe.Label",
                        font=("Segoe UI", 10, "bold"),
                        foreground="#4a90e2",
                        background="#ffffff")

        self.in_var = tk.StringVar()
        self.out_var = tk.StringVar(value="")
        self.lang = tk.StringVar(value="en")
        self.dpi = tk.IntVar(value=150)
        self.close_chrome = tk.BooleanVar(value=True)

        # Sidebar
        sidebar = ttk.Frame(self, style="Sidebar.TFrame", padding=20)
        sidebar.pack(side="left", fill="y")

        ttk.Label(sidebar, text="📂 Input PDF", foreground="#2c3e50", background="#e8ebef").pack(anchor="w", pady=(0, 4))
        ttk.Entry(sidebar, textvariable=self.in_var, width=32).pack(pady=(0, 10))
        ttk.Button(sidebar, text="Browse", command=self.pick_in).pack(pady=(0, 16), fill="x")

        ttk.Label(sidebar, text="💾 Output PDF", foreground="#2c3e50", background="#e8ebef").pack(anchor="w", pady=(0, 4))
        ttk.Entry(sidebar, textvariable=self.out_var, width=32).pack(pady=(0, 10))
        ttk.Button(sidebar, text="Choose", command=self.pick_out).pack(pady=(0, 16), fill="x")

        # Main area
        main = ttk.Frame(self, style="Card.TFrame", padding=20)
        main.pack(side="left", fill="both", expand=True)

        # Options card
        opt = ttk.Labelframe(main, text="Options", padding=12, style="Card.TFrame")
        opt.pack(fill="x", pady=(0, 12))
        ttk.Label(opt, text="🌍 Target language:").grid(row=0, column=0, sticky="w")
        ttk.Entry(opt, textvariable=self.lang, width=10).grid(row=0, column=1, padx=6)
        ttk.Label(opt, text="🖼 DPI:").grid(row=0, column=2, sticky="w")
        ttk.Entry(opt, textvariable=self.dpi, width=8).grid(row=0, column=3, padx=6)
        ttk.Checkbutton(opt, text="Close Chrome after translation", variable=self.close_chrome).grid(row=0, column=4, padx=12)

    
        # Log card
        log_card = ttk.Labelframe(main, text="Logs", padding=12, style="Card.TFrame")
        log_card.pack(fill="both", expand=True, pady=(0, 12))
        self.log = tk.Text(log_card,
                           height=12,
                           font=("Segoe UI", 10),
                           wrap="word",
                           bg="#fafafa",
                           fg="#333",
                           highlightthickness=1,
                           highlightbackground="#dcdfe3",
                           relief="flat",
                           padx=10,
                           pady=6)
        self.log.pack(fill="both", expand=True)

        # Buttons
        btns = ttk.Frame(main, style="Card.TFrame")
        btns.pack(fill="x", pady=(0, 5))
        self.run_btn = ttk.Button(btns, text="🚀 Translate PDF", command=self.start)
        self.run_btn.pack(side="left", padx=6)
        ttk.Button(btns, text="❌ Quit", command=self.destroy).pack(side="right", padx=6)

    def pick_in(self):
        f = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
        if f: self.in_var.set(f)

    def pick_out(self):
        f = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])
        if f: self.out_var.set(f)

    def log_write(self, msg):
        self.log.insert("end", msg + "\n")
        self.log.see("end")
        self.update_idletasks()

    def start(self):
        if not self.in_var.get():
            messagebox.showwarning(APP_NAME, "Please choose an input PDF.")
            return
        if not self.out_var.get():
            messagebox.showwarning(APP_NAME, "Please choose where to save the output.")
            return
        self.run_btn.configure(state="disabled")
        def worker():
            try:
                asyncio.run(translate_pdf(
                    self.in_var.get(), self.out_var.get(),
                    target_lang=self.lang.get().strip() or "en",
                    dpi=int(self.dpi.get()),
                    close_browser=bool(self.close_chrome.get()),
                    log=self.log_write
                ))
                self.log_write("✔ Translation complete! Your shiny new PDF + DOCX are ready.")
            except Exception as e:
                self.log_write(f"❌ Error: {e}")
                messagebox.showerror(APP_NAME, str(e))
            finally:
                self.progress.stop()
                self.run_btn.configure(state="normal")
        Thread(target=worker, daemon=True).start()



if __name__ == "__main__":
    App().mainloop()
