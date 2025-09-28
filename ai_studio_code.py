import tkinter as tk
from tkinter import filedialog, scrolledtext, ttk, messagebox
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
from datetime import datetime
import pythoncom

# Attempt to import Windows-specific libraries for TOC updates
try:
    import win32com.client
    IS_WINDOWS = True
except ImportError:
    IS_WINDOWS = False

# =====================================================================
#  ENHANCED DOCX PROCESSING AND TABLE DETECTION LOGIC
# =====================================================================

class EnhancedDocxProcessor:
    """Handles advanced DOCX operations like TOC updates."""

    @staticmethod
    def update_toc_with_com(filepath):
        """
        Updates the Table of Contents using Word's COM automation on Windows.
        This respects the existing formatting and styles of the TOC.
        """
        if not IS_WINDOWS:
            messagebox.showwarning("Unsupported OS", "TOC update feature is only available on Windows with MS Word installed.")
            return False
        try:
            pythoncom.CoInitialize()
            word_app = win32com.client.Dispatch("Word.Application")
            doc = word_app.Documents.Open(os.path.abspath(filepath))
            doc.TablesOfContents(1).Update()
            doc.Close(SaveChanges=True)
            word_app.Quit()
            pythoncom.CoUninitialize()
            return True
        except Exception as e:
            print(f"Error updating TOC with COM: {e}")
            try:
                pythoncom.CoUninitialize()
            except pythoncom.error:
                pass
            return False

    @staticmethod
    def add_toc_to_document(doc):
        """
        Adds a placeholder for a Table of Contents if one is missing.
        """
        doc.paragraphs[0].insert_paragraph_before("Table of Contents")
        run = doc.paragraphs[0].insert_paragraph_before().add_run()
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar_begin)
        run._r.append(instrText)
        run._r.append(fldChar_end)


class TableRuleEngine:
    """
    Detects table types and provides specific formatting rules based on new visual evidence.
    """
    @staticmethod
    def get_table_config(table, preceding_paragraph_text=""):
        """Analyzes a table and its context, returning a dictionary of formatting rules."""
        headers = [cell.text.lower().strip() for cell in table.rows[0].cells] if table.rows else []
        header_text = " ".join(headers)
        
        first_row_content = [cell.text.lower().strip() for cell in table.rows[0].cells] if table.rows else []

        # Rule for Reference List Table (Image 2) - Left Aligned
        reference_list_headers = ['publication number', 'priority date', 'filing date', 'inventor(s)', 'assignee(s)']
        if any(h in header_text for h in reference_list_headers):
            return {
                'type': 'Reference List',
                'header_align': WD_ALIGN_PARAGRAPH.CENTER,
                'body_align': WD_ALIGN_PARAGRAPH.CENTER
            }

        # Rule for Legend/Definition Table (Image 3) - Mixed Alignment
        legend_keywords = ['supported:', 'inferentially supported:', 'partially supported:', 'not supported:']
        if any(keyword in content for content in first_row_content for keyword in legend_keywords):
            return {
                'type': 'Legend Table',
                'header_align': WD_ALIGN_PARAGRAPH.CENTER,
                'body_align': WD_ALIGN_PARAGRAPH.JUSTIFY, # Default, will be overridden
                'special_rules': {'legend_formatting': True}
            }

        # Rule for Claim Chart / Claim Matrix (Image 4) - Justified with centered symbols
        if 'claim element' in header_text:
            return {
                'type': 'Claim Chart',
                'header_align': WD_ALIGN_PARAGRAPH.CENTER,
                'body_align': WD_ALIGN_PARAGRAPH.JUSTIFY
            }

        # Default rule for all other tables (e.g., Bibliographic Data from Image 1) is JUSTIFIED
        return {
            'type': 'Standard Table',
            'header_align': WD_ALIGN_PARAGRAPH.CENTER,
            'body_align': WD_ALIGN_PARAGRAPH.JUSTIFY
        }


# =====================================================================
#  MAIN TKINTER APPLICATION
# =====================================================================

class DocxFormatChecker:
    def __init__(self, root):
        self.root = root
        self.root.title("Advanced DOCX Report Reviewer")
        self.root.geometry("600x400")
        
        self.style = ttk.Style()
        self.style.theme_use('clam')
        
        self.current_file = None
        self.doc = None
        self.issues = []
        
        self.setup_ui()
        
    def setup_ui(self):
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky="nsew")
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(2, weight=1)
        
        title_label = ttk.Label(main_frame, text="Advanced DOCX Report Reviewer", font=('Arial', 16, 'bold'))
        title_label.grid(row=0, column=0, columnspan=2, pady=(0, 10))
        
        file_frame = ttk.LabelFrame(main_frame, text="File Operations", padding=10)
        file_frame.grid(row=1, column=0, columnspan=2, sticky="ew", pady=5)
        self.file_label = ttk.Label(file_frame, text="No file selected.")
        self.file_label.pack(side=tk.LEFT, padx=5)
        
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=1, column=1, sticky="e")
        ttk.Button(button_frame, text="Browse...", command=self.browse_file).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="1. Check Document", command=self.check_format).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="2. Apply All Fixes", command=self.apply_fixes).pack(side=tk.LEFT, padx=5)
        self.update_toc_button = ttk.Button(button_frame, text="3. Update TOC", command=self.update_toc)
        self.update_toc_button.pack(side=tk.LEFT, padx=5)
        if not IS_WINDOWS:
            self.update_toc_button.config(state=tk.DISABLED)

        results_frame = ttk.LabelFrame(main_frame, text="Check Results", padding=5)
        results_frame.grid(row=2, column=0, columnspan=2, sticky="nsew", pady=5)
        results_frame.columnconfigure(0, weight=1)
        results_frame.rowconfigure(0, weight=1)
        
        self.results_text = scrolledtext.ScrolledText(results_frame, wrap=tk.WORD, font=('Consolas', 10))
        self.results_text.grid(row=0, column=0, sticky="nsew")
        self.setup_tags()
        
        self.status_label = ttk.Label(main_frame, text="Ready", relief=tk.SUNKEN)
        self.status_label.grid(row=4, column=0, columnspan=2, sticky="ew")

    def setup_tags(self):
        self.results_text.tag_config("error", foreground="#D32F2F", font=('Consolas', 10, 'bold'))
        self.results_text.tag_config("success", foreground="#388E3C", font=('Consolas', 10, 'bold'))
        self.results_text.tag_config("info", foreground="#1976D2")
        self.results_text.tag_config("header", font=('Consolas', 12, 'bold'), foreground="#004D40")

    def browse_file(self):
        filename = filedialog.askopenfilename(title="Select DOCX file", filetypes=[("Word Documents", "*.docx")])
        if filename:
            self.current_file = filename
            self.file_label.config(text=os.path.basename(filename))
            self.status_label.config(text=f"Loaded: {os.path.basename(filename)}")
            self.results_text.delete(1.0, tk.END)

    def check_format(self):
        if not self.current_file:
            messagebox.showwarning("No File", "Please select a DOCX file first.")
            return
        
        self.issues = []
        self.results_text.delete(1.0, tk.END)
        self.status_label.config(text="Checking document...")
        self.root.update_idletasks()

        try:
            self.doc = Document(self.current_file)
            self.check_body_text()
            self.check_headings()
            self.check_tables()
            self.check_images()
            self.check_line_spacing()
            self.check_document_properties()
            self.check_toc()
            self.check_toc_font()
            self.display_summary()
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred while checking the document:\n{str(e)}")
            self.status_label.config(text="Error occurred.")

    def display_summary(self):
        self.results_text.insert(tk.END, "\n" + "="*80 + "\n", "header")
        self.results_text.insert(tk.END, "CHECK SUMMARY\n", "header")
        if not self.issues:
            self.results_text.insert(tk.END, "✓ Congratulations! All quality checks passed.\n", "success")
        else:
            self.results_text.insert(tk.END, f"Found {len(self.issues)} issues to address.\n", "error")
            self.results_text.insert(tk.END, "Click 'Apply All Fixes' to correct them automatically.\n", "info")
        self.status_label.config(text="Check complete.")
        
    # --- Individual Check Methods ---

    def check_body_text(self):
        self.results_text.insert(tk.END, "1. Body Text\n", "header")
        font_exceptions = ['Segoe UI', 'Wingdings', 'Wingdings 2']
        found = False
        for i, para in enumerate(self.doc.paragraphs):
            is_image_para = bool(para._element.xpath('.//w:drawing'))
            
            if para.style.name == 'Normal' and para.text.strip():
                if not is_image_para and para.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
                    self.issues.append({'type': 'body_alignment', 'paragraph': i})
                    self.results_text.insert(tk.END, f"  ✗ Body text (Para {i+1}): Not justified.\n", "error")
                    found = True
                    
                for run in para.runs:
                    if run.font.name not in font_exceptions and run.font.name != 'Calibri':
                        self.issues.append({'type': 'body_font', 'paragraph': i})
                        self.results_text.insert(tk.END, f"  ✗ Body text (Para {i+1}): Incorrect font '{run.font.name}'.\n", "error")
                        found = True
                        break
        if not found: self.results_text.insert(tk.END, "  ✓ Correct\n", "success")

    def check_headings(self):
        self.results_text.insert(tk.END, "2. Headings\n", "header")
        specs = {'Heading 1': 28, 'Heading 2': 20, 'Heading 3': 14}
        found = False
        for i, p in enumerate(self.doc.paragraphs):
            if p.style.name in specs:
                if p.alignment != WD_ALIGN_PARAGRAPH.LEFT:
                    self.issues.append({'type': 'heading_format', 'paragraph': i})
                    self.results_text.insert(tk.END, f"  ✗ {p.style.name} (Para {i+1}): Not left-aligned.\n", "error")
                    found = True
                for run in p.runs:
                    if run.font.name != 'Cambria' or run.font.size != Pt(specs[p.style.name]):
                        self.issues.append({'type': 'heading_format', 'paragraph': i})
                        self.results_text.insert(tk.END, f"  ✗ {p.style.name} (Para {i+1}): Incorrect font or size.\n", "error")
                        found = True
                        break
        if not found: self.results_text.insert(tk.END, "  ✓ Correct\n", "success")

    def check_tables(self):
        self.results_text.insert(tk.END, "3. Tables\n", "header")
        if not self.doc.tables:
            self.results_text.insert(tk.END, "  - No tables found.\n", "info")
            return
        
        self.issues.append({'type': 'fix_all_tables'})
        self.results_text.insert(tk.END, "  - All tables scheduled for formatting review and fix.\n", "info")
        
    def check_images(self):
        self.results_text.insert(tk.END, "4. Images\n", "header")
        found = False
        for i, p in enumerate(self.doc.paragraphs):
            if p._element.xpath('.//w:drawing') and p.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                self.issues.append({'type': 'image_alignment', 'paragraph': i})
                self.results_text.insert(tk.END, f"  ✗ Image near paragraph {i+1}: Not centered.\n", "error")
                found = True
        if not found: self.results_text.insert(tk.END, "  ✓ Correct\n", "success")

    def check_line_spacing(self):
        self.results_text.insert(tk.END, "5. Spacing After Heading 1\n", "header")
        found = False
        in_h1_content = False
        for i, p in enumerate(self.doc.paragraphs):
            if p.style.name == 'Heading 1':
                in_h1_content = True
                continue
            if p.style.name in ['Heading 2', 'Heading 3']:
                in_h1_content = False
            if in_h1_content and p.text.strip():
                pf = p.paragraph_format
                if pf.space_before != Pt(6) or pf.space_after != Pt(6) or (pf.line_spacing is not None and abs(pf.line_spacing - 1.33) > 0.01):
                    self.issues.append({'type': 'line_spacing', 'paragraph': i})
                    self.results_text.insert(tk.END, f"  ✗ Content at paragraph {i+1}: Incorrect line/paragraph spacing.\n", "error")
                    found = True
        if not found: self.results_text.insert(tk.END, "  ✓ Correct\n", "success")

    def check_document_properties(self):
        self.results_text.insert(tk.END, "6. Document Properties\n", "header")
        base_name = os.path.splitext(os.path.basename(self.current_file))[0]
        props = self.doc.core_properties
        if any(getattr(props, p) != base_name for p in ['title', 'subject', 'keywords', 'category', 'comments']):
            self.issues.append({'type': 'doc_properties'})
            self.results_text.insert(tk.END, "  ✗ Properties do not match filename.\n", "error")
        else:
            self.results_text.insert(tk.END, "  ✓ Correct\n", "success")

    def check_toc(self):
        self.results_text.insert(tk.END, "7. Table of Contents\n", "header")
        if not any('TOC' in p._element.xml for p in self.doc.paragraphs):
            self.issues.append({'type': 'toc_missing'})
            self.results_text.insert(tk.END, "  ✗ TOC is missing from the document.\n", "error")
        else:
            self.results_text.insert(tk.END, "  ✓ TOC is present.\n", "success")

    def check_toc_font(self):
        self.results_text.insert(tk.END, "8. TOC Font\n", "header")
        found_issue = False
        is_present = any('TOC' in p._element.xml for p in self.doc.paragraphs)
        if not is_present:
            self.results_text.insert(tk.END, "  - TOC not present, skipping font check.\n", "info")
            return

        for p in self.doc.paragraphs:
            if p.style and p.style.name.startswith('TOC'):
                for run in p.runs:
                    if run.font.name != 'Calibri' or run.font.size != Pt(11):
                        self.issues.append({'type': 'toc_font'})
                        self.results_text.insert(tk.END, f"  ✗ TOC Style '{p.style.name}': Incorrect font ('{run.font.name}') or size.\n", "error")
                        found_issue = True
                        break 
                if found_issue:
                    break 
        if not found_issue:
            self.results_text.insert(tk.END, "  ✓ Correct\n", "success")
            
    # --- Fixing Logic ---

    def apply_fixes(self):
        if not self.issues:
            messagebox.showinfo("No Issues", "There are no issues to fix.")
            return

        try:
            backup_file = self.current_file.replace('.docx', f'_backup_{datetime.now():%Y%m%d%H%M%S}.docx')
            self.doc.save(backup_file)

            applied_fixes = set()

            for issue in self.issues:
                fix_type = issue['type']
                if fix_type not in applied_fixes or 'paragraph' in issue:
                    self.fix_issue(issue)
                    if 'paragraph' not in issue:
                        applied_fixes.add(fix_type)

            fixed_file = self.current_file.replace('.docx', '_fixed.docx')
            self.doc.save(fixed_file)
            messagebox.showinfo("Success", f"Fixes applied successfully.\n\nA backup was saved as:\n{os.path.basename(backup_file)}\n\nThe corrected file is:\n{os.path.basename(fixed_file)}")
            self.status_label.config(text="Fixes applied.")
        except Exception as e:
            messagebox.showerror("Error Applying Fixes", f"An error occurred: {str(e)}")

    def fix_issue(self, issue):
        fix_type = issue['type']
        font_exceptions = ['Wingdings', 'Wingdings 2']
        
        if fix_type == 'body_alignment':
            self.doc.paragraphs[issue['paragraph']].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif fix_type == 'body_font':
            p = self.doc.paragraphs[issue['paragraph']]
            for run in p.runs:
                if run.font.name not in font_exceptions:
                    run.font.name = 'Segoe UI'
                    run.font.size = Pt(10)
        
        elif fix_type == 'heading_format':
            p = self.doc.paragraphs[issue['paragraph']]
            specs = {'Heading 1': 28, 'Heading 2': 20, 'Heading 3': 14}
            size = specs.get(p.style.name)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Cambria'
                if size: run.font.size = Pt(size)
        
        elif fix_type == 'image_alignment':
            self.doc.paragraphs[issue['paragraph']].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        elif fix_type == 'line_spacing':
            p = self.doc.paragraphs[issue['paragraph']]
            pf = p.paragraph_format
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            pf.line_spacing = 1.33
        
        elif fix_type == 'doc_properties':
            base_name = os.path.splitext(os.path.basename(self.current_file))[0]
            props = self.doc.core_properties
            for p_name in ['title', 'subject', 'keywords', 'category', 'comments']:
                setattr(props, p_name, base_name)
        
        elif fix_type == 'toc_missing':
            EnhancedDocxProcessor.add_toc_to_document(self.doc)
        
        elif fix_type == 'toc_font':
            self.format_toc_font(self.doc)

        elif fix_type == 'fix_all_tables':
            self.format_all_tables()

    def get_paragraph_before(self, element):
        """Finds the paragraph element immediately before the given element (e.g., a table)."""
        try:
            body_children = list(self.doc.element.body)
            element_idx = body_children.index(element._element)
            if element_idx > 0:
                prev_elm = body_children[element_idx - 1]
                if prev_elm.tag.endswith('p'):
                    return docx.text.paragraph.Paragraph(prev_elm, self.doc)
        except (ValueError, IndexError):
            pass
        return None

    def format_all_tables(self):
        for table in self.doc.tables:
            prev_para = self.get_paragraph_before(table)
            prev_para_text = prev_para.text if prev_para else ""
            self.format_table(table, prev_para_text)
            
    def format_toc_font(self, doc):
        """Sets the font for paragraphs with TOC styles to Calibri 11pt."""
        for p in doc.paragraphs:
            if p.style and p.style.name.startswith('TOC'):
                for run in p.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)

    def format_table(self, table, preceding_text=""):
        config = TableRuleEngine.get_table_config(table, preceding_text)
        font_exceptions = ['Wingdings', 'Wingdings 2']
        symbols = ['✓', '☒', 'P', '-']

        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                is_header_row = (r_idx == 0)

                for p in cell.paragraphs:
                    # Determine alignment based on new rules
                    alignment = config.get('body_align', WD_ALIGN_PARAGRAPH.JUSTIFY)

                    if is_header_row:
                        alignment = config['header_align']
                    
                    # Rule for Legend Table (Image 3)
                    elif config.get('special_rules', {}).get('legend_formatting'):
                        alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 else WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    # Rule for Claim Charts (Image 4): Center symbols and # column
                    elif config['type'] == 'Claim Chart':
                        # Center align if it's the first column (e.g., '#') or if the cell text is just a symbol
                        if c_idx == 0 or p.text.strip() in symbols:
                            alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    p.alignment = alignment

                    # Format fonts
                    for run in p.runs:
                        if run.font.name not in font_exceptions:
                            run.font.name = 'Segoe UI'
                            run.font.size = Pt(10)

    def update_toc(self):
        if not self.current_file:
            messagebox.showwarning("No File", "Please load a file first.")
            return

        fixed_file = self.current_file.replace('.docx', '_fixed.docx')
        target_file = fixed_file if os.path.exists(fixed_file) else self.current_file

        self.status_label.config(text="Updating TOC... This requires MS Word and may take a moment.")
        self.root.update_idletasks()

        if EnhancedDocxProcessor.update_toc_with_com(target_file):
            messagebox.showinfo("Success", f"Table of Contents updated in:\n{os.path.basename(target_file)}")
            self.status_label.config(text="TOC Updated.")
        else:
            messagebox.showerror("TOC Update Failed", "Could not update the TOC. Please ensure MS Word is installed and not busy.")
            self.status_label.config(text="TOC update failed.")

def main():
    root = tk.Tk()
    app = DocxFormatChecker(root)
    root.mainloop()

if __name__ == "__main__":
    main()